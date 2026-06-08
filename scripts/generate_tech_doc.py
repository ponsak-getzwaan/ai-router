"""Generate Technical Architecture Document for AI Router.

Produces reports/AI_Router_Technical_Architecture.docx with colour-coded
block diagrams built from python-docx tables (no external image dependencies).

Usage:
    uv run python scripts/generate_tech_doc.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Brand colours
# ---------------------------------------------------------------------------
BRAND_DARK  = RGBColor(0x0E, 0x3B, 0x2A)
BRAND_MID   = RGBColor(0x5A, 0x6B, 0x64)
BRAND_LIGHT = RGBColor(0xF5, 0xF2, 0xEB)
TEXT_DARK   = RGBColor(0x1A, 0x1F, 0x2E)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

# Layer accent colours (match admin dashboard dots)
C_ORCH       = "0E3B2A"   # dark green  — Orchestrator
C_BOUNCER    = "1D4ED8"   # blue        — Bouncer
C_CLASSIFIER = "7C3AED"   # purple      — Classifier
C_STRATEGIST = "D97706"   # amber       — Strategist
C_ADAPTER    = "475569"   # slate grey  — Adapter/Vendor
C_REDACTOR   = "15803D"   # green       — Redactor/PII
C_AUTH       = "0891B2"   # teal        — Auth
C_INFRA      = "374151"   # dark grey   — Infrastructure
C_HEADER     = "0E3B2A"   # brand dark  — Table headers
C_ALT        = "F0FDF4"   # very light green — alternating rows

# Hex shades for cell backgrounds (lighter versions)
SH_ORCH       = "D1FAE5"
SH_BOUNCER    = "DBEAFE"
SH_CLASSIFIER = "EDE9FE"
SH_STRATEGIST = "FEF3C7"
SH_ADAPTER    = "F1F5F9"
SH_REDACTOR   = "DCFCE7"
SH_AUTH       = "E0F2FE"
SH_INFRA      = "F3F4F6"
SH_HEADER     = "0E3B2A"
SH_ALT        = "F9FAFB"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _shade_cell(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_border(cell, top="none", bottom="none", left="none", right="none",
                     color="CCCCCC", size="4") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), val)
        if val != "none":
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _cell_para(cell, text: str, bold: bool = False, italic: bool = False,
               size: int = 9, color: RGBColor | None = None,
               align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def _add_cell_line(cell, text: str, bold: bool = False, italic: bool = False,
                   size: int = 9, color: RGBColor | None = None) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_toc(doc: Document) -> None:
    para = doc.add_paragraph()
    para.style = "Normal"
    run = para.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2 = para.add_run()
    run2._r.append(instr)
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    run3 = para.add_run()
    run3._r.append(fld2)
    placeholder = doc.add_paragraph(
        "[ Right-click this line in Word and choose Update Field to generate the Table of Contents ]"
    )
    placeholder.style = "Normal"
    placeholder.runs[0].font.color.rgb = BRAND_MID
    placeholder.runs[0].font.italic = True
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    run4 = para.add_run()
    run4._r.append(fld3)


def _set_styles(doc: Document) -> None:
    s = doc.styles
    h1 = s["Heading 1"]
    h1.font.name = "Calibri"; h1.font.size = Pt(18); h1.font.color.rgb = BRAND_DARK; h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18); h1.paragraph_format.space_after = Pt(6)

    h2 = s["Heading 2"]
    h2.font.name = "Calibri"; h2.font.size = Pt(13); h2.font.color.rgb = BRAND_DARK; h2.font.bold = True
    h2.paragraph_format.space_before = Pt(14); h2.paragraph_format.space_after = Pt(4)

    h3 = s["Heading 3"]
    h3.font.name = "Calibri"; h3.font.size = Pt(11); h3.font.color.rgb = TEXT_DARK; h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10); h3.paragraph_format.space_after = Pt(3)

    n = s["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(10); n.font.color.rgb = TEXT_DARK
    n.paragraph_format.space_after = Pt(5)
    n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    n.paragraph_format.line_spacing = 1.15


# ---------------------------------------------------------------------------
# Content helpers
# ---------------------------------------------------------------------------

def H(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def P(doc: Document, text: str, italic: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph(text, style="Normal")
    if italic or size != 10:
        for r in p.runs:
            r.italic = italic
            r.font.size = Pt(size)


def note(doc: Document, label: str, text: str, accent: str = "1D4ED8") -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.cell(0, 0)
    _shade_cell(c, "EFF6FF")
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(
        int(accent[0:2], 16), int(accent[2:4], 16), int(accent[4:6], 16)
    )
    r2 = p.add_run(text)
    r2.font.size = Pt(9)
    r2.font.color.rgb = TEXT_DARK
    doc.add_paragraph()


def warn(doc: Document, label: str, text: str) -> None:
    note(doc, label, text, accent="D97706")


def danger(doc: Document, label: str, text: str) -> None:
    note(doc, label, text, accent="DC2626")


def B(doc: Document, label: str, desc: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(label)
    r1.bold = True
    if desc:
        r2 = p.add_run(f"  —  {desc}")
        r2.font.color.rgb = BRAND_MID
        r2.font.size = Pt(9)


def arrow_down(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("v")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = BRAND_MID


def diagram_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = BRAND_MID


# ---------------------------------------------------------------------------
# Diagram: Pipeline Overview (vertical coloured blocks)
# ---------------------------------------------------------------------------

def _flow_block(doc: Document, fill: str, title: str, lines: list[str],
                text_color: str = "FFFFFF") -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    c = t.cell(0, 0)
    _shade_cell(c, fill)
    tc = c.paragraphs[0]
    tc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc.paragraph_format.space_before = Pt(5)
    tc.paragraph_format.space_after = Pt(2)
    rgb = RGBColor(int(text_color[0:2], 16), int(text_color[2:4], 16), int(text_color[4:6], 16))
    r = tc.add_run(title)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = rgb
    for line in lines:
        lp = c.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(1)
        lr = lp.add_run(line)
        lr.font.size = Pt(8)
        lr.font.color.rgb = rgb


def diagram_pipeline_overview(doc: Document) -> None:
    blocks = [
        ("1D4ED8", "CLIENT  &  API GATEWAY",
         ["HTTPS POST /chat  |  Bearer JWT", "Cognito Lambda Authoriser validates token at edge",
          "202 Accepted + correlation_id returned immediately"], "FFFFFF"),
        ("0891B2", "SQS INCOMING QUEUE",
         ["Decouples ingest from processing", "Message: user_sub, raw message, source_ip, correlation_id"], "FFFFFF"),
        ("0E3B2A", "ORCHESTRATOR",
         ["Set correlation_id ContextVar", "Call Presidio sidecar  ->  tokenise PII",
          "Store tokens in Redis vault (TTL 5 min)", "Build immutable PipelineEnvelope (redacted only)"], "FFFFFF"),
        ("15803D", "REDACTOR  (Presidio Sidecar)",
         ["Detects SG_NRIC, SG_FIN, EMAIL, CREDIT_CARD ...",
          "Replaces each entity with VAULT:{corr_id}:{token}",
          "Raw message stays inside Orchestrator — never propagated"], "FFFFFF"),
        ("1D4ED8", "BOUNCER  [200 ms total budget]",
         ["Stage 1: Rule Gate  (microseconds) — length, regex, banned user, rate limit",
          "Stage 2: Haiku micro-classifier  (remaining budget)",
          "PASS  /  BLOCK  /  ESCALATE  /  FAIL-OPEN on timeout"], "FFFFFF"),
        ("7C3AED", "INTENT CLASSIFIER",
         ["Fast Path: Cohere embedding similarity  (~20 ms)",
          "Deep Path: Claude Sonnet + MCP tools  (complex queries)",
          "Output: ClassifiedIntent  (intent, confidence, resolved_message)"], "FFFFFF"),
        ("D97706", "ROUTING STRATEGIST",
         ["Confidence >= 0.85  ->  deterministic rule lookup",
          "0.50 - 0.85  ->  Haiku arbitration call",
          "< 0.50  ->  escalate to human review",
          "Policy engine: data residency, compliance blocks"], "FFFFFF"),
        ("475569", "VENDOR ADAPTER  (LiteLLM + Bedrock)",
         ["Routes to Bedrock inference profile (APAC/Global)",
          "Streaming response via InvokeModelWithResponseStream",
          "Supported: Claude Sonnet, Claude Haiku, Llama, Mistral"], "FFFFFF"),
        ("0E3B2A", "OUTPUT HANDLER  +  AUDIT",
         ["Per-chunk: vault restore  ->  leak detector  ->  WebSocket/SSE",
          "Finally block: DynamoDB audit write, vault cleanup",
          "Audit stores entity types & counts — never raw values or PII"], "FFFFFF"),
    ]
    for fill, title, lines, tc in blocks:
        _flow_block(doc, fill, title, lines, tc)
        arrow_down(doc)
    # Remove last arrow
    doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)


# ---------------------------------------------------------------------------
# Diagram: Bouncer Two-Stage Gate
# ---------------------------------------------------------------------------

def diagram_bouncer(doc: Document) -> None:
    t = doc.add_table(rows=7, cols=3)
    t.style = "Table Grid"
    widths = [Inches(1.4), Inches(2.8), Inches(2.0)]

    # Header row
    hdr_labels = ["Stage", "Logic", "Outcome"]
    for i, label in enumerate(hdr_labels):
        c = t.cell(0, i)
        _shade_cell(c, C_HEADER)
        _cell_para(c, label, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    rows_data = [
        (SH_BOUNCER, "Stage 1\nRule Gate", SH_BOUNCER,
         "Length check (max 4,096 chars)\nPrompt injection regex patterns\nBanned user_sub lookup (Redis)\nRate limit check (Redis)",
         SH_BOUNCER, "REJECT -> BounceResult(allowed=False)\nPASS -> proceed to Stage 2"),
        (SH_BOUNCER, "", "F0FDF4",
         "Stage 2   Haiku Micro-Classifier\nClaude Haiku via Bedrock (max_tokens=50)\nBudget: remaining ms of 200ms total",
         SH_BOUNCER, ""),
        (SH_BOUNCER, "confidence >= 0.7\nAND pass = true", "D1FAE5",
         "-> BounceResult(allowed=True, layer='llm_classifier')", SH_BOUNCER, "ALLOW downstream"),
        (SH_BOUNCER, "confidence < 0.7", "FEF9C3",
         "-> BounceResult(allowed=True, escalate=True)\n+ async DynamoDB review log entry", SH_BOUNCER, "ESCALATE (but allow)"),
        (SH_BOUNCER, "200ms budget\nexceeded", "FEE2E2",
         "FAIL-OPEN: BounceResult(allowed=True, timed_out=True)\nCloudWatch: BouncerTimeout += 1", SH_BOUNCER, "ALLOW (fail-open)"),
        (SH_BOUNCER, "Bedrock error", "FEE2E2",
         "FAIL-OPEN: BounceResult(allowed=True, timed_out=False)\nCloudWatch: BouncerError += 1", SH_BOUNCER, "ALLOW (fail-open)"),
    ]

    for r_idx, (f0, t0, f1, t1, f2, t2) in enumerate(rows_data, start=1):
        row = t.rows[r_idx]
        c0, c1, c2 = row.cells[0], row.cells[1], row.cells[2]
        _shade_cell(c0, f0); _shade_cell(c1, f1); _shade_cell(c2, f2)
        _cell_para(c0, t0, bold=True, size=8, color=TEXT_DARK)
        _cell_para(c1, t1, size=8, color=TEXT_DARK)
        _cell_para(c2, t2, bold=(r_idx in (3, 4, 5, 6)), size=8, color=TEXT_DARK)


# ---------------------------------------------------------------------------
# Diagram: Classifier Two Paths
# ---------------------------------------------------------------------------

def diagram_classifier(doc: Document) -> None:
    t = doc.add_table(rows=6, cols=2)
    t.style = "Table Grid"

    hdr = t.rows[0].cells
    for i, label in enumerate(["FAST PATH  (embedding similarity)", "DEEP PATH  (Sonnet + MCP)"]):
        fill = "7C3AED" if i == 0 else "4C1D95"
        _shade_cell(hdr[i], fill)
        _cell_para(hdr[i], label, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    fast_rows = [
        ("When used", "Short, simple messages\nClear intent signals\nMatches known exemplars"),
        ("Model", "Cohere Embed Multilingual v3\n(Bedrock, search_query input type)"),
        ("Mechanism", "Cosine similarity vs pre-loaded\nexemplar vectors per intent domain"),
        ("Threshold", "Similarity >= 0.70  ->  return result\n< 0.70  ->  fall through to Deep Path"),
        ("Latency", "~20 ms  (embedding only, no LLM call)"),
    ]
    deep_rows = [
        ("When used", "Complex / multi-domain queries\nFast path score below threshold\nPronoun resolution needed"),
        ("Model", "Claude Sonnet 4.6 via Bedrock\n(global inference profile)"),
        ("MCP Tools", "get_intent_taxonomy  (DynamoDB)\nget_session_history  (Redis)\ncheck_guardrails  (Bedrock Guardrails)\nget_entity_context  (Redis)"),
        ("Threshold", "Confidence >= 0.60  ->  return result\n< 0.60  ->  escalate to human review"),
        ("Latency", "2 – 8 s  (full Sonnet call + MCP round-trips)"),
    ]

    for r_idx in range(1, 6):
        row = t.rows[r_idx]
        fill_f = SH_CLASSIFIER if r_idx % 2 == 1 else "F5F3FF"
        fill_d = SH_CLASSIFIER if r_idx % 2 == 1 else "F5F3FF"
        _shade_cell(row.cells[0], fill_f)
        _shade_cell(row.cells[1], fill_d)
        label_f, text_f = fast_rows[r_idx - 1]
        label_d, text_d = deep_rows[r_idx - 1]
        _cell_para(row.cells[0], label_f + "\n" + text_f, size=8, color=TEXT_DARK)
        _cell_para(row.cells[1], label_d + "\n" + text_d, size=8, color=TEXT_DARK)
        # Bold the label part
        for cell, label in [(row.cells[0], label_f), (row.cells[1], label_d)]:
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].bold = True


# ---------------------------------------------------------------------------
# Diagram: Routing Strategist Confidence Tiers
# ---------------------------------------------------------------------------

def diagram_strategist(doc: Document) -> None:
    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"

    for i, label in enumerate(["Confidence Tier", "Action", "Details"]):
        c = t.cell(0, i)
        _shade_cell(c, C_HEADER)
        _cell_para(c, label, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    tiers = [
        ("D1FAE5", ">= 0.85\nHIGH CONFIDENCE", "D1FAE5",
         "Deterministic rule lookup\n(no LLM call)", "D1FAE5",
         "Read routing rule from DynamoDB cache\nCheck vendor health (parallel)\nApply policy engine\nReturn RoutingPlan immediately"),
        ("FEF3C7", "0.50 - 0.85\nMEDIUM CONFIDENCE", "FEF3C7",
         "Haiku arbitration call\n(max 80 tokens)", "FEF3C7",
         "Haiku weighs candidate vendors\nConsiders session history & intent\nFallback chain selected per intent\nPolicy engine runs after selection"),
        ("FEE2E2", "< 0.50\nLOW CONFIDENCE", "FEE2E2",
         "Escalate to human review\n(Review Log + Degradation Ladder)", "FEE2E2",
         "Hedge -> Clarify -> Fallback -> Hand-off\nAsync DynamoDB review log entry\nUser gets graceful degradation response"),
    ]

    for r_idx, (f0, t0, f1, t1, f2, t2) in enumerate(tiers, start=1):
        row = t.rows[r_idx]
        _shade_cell(row.cells[0], f0)
        _shade_cell(row.cells[1], f1)
        _shade_cell(row.cells[2], f2)
        _cell_para(row.cells[0], t0, bold=True, size=8, color=TEXT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[1], t1, bold=True, size=8, color=TEXT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[2], t2, size=8, color=TEXT_DARK)


# ---------------------------------------------------------------------------
# Diagram: PII Redaction Lifecycle
# ---------------------------------------------------------------------------

def diagram_redaction(doc: Document) -> None:
    stages = [
        ("15803D", "1. INGEST\n(Orchestrator)", "FFFFFF",
         "Raw message\nreceived"),
        ("15803D", "2. DETECT\n(Presidio sidecar)", "FFFFFF",
         "Recogniser scans\nfor PII entities"),
        ("15803D", "3. TOKENISE\n(Orchestrator)", "FFFFFF",
         "Replace entity with\nVAULT token"),
        ("15803D", "4. STORE\n(Redis vault)", "FFFFFF",
         "vault:{corr}:{token}\nTTL: 5 minutes"),
        ("15803D", "5. PROCESS\n(LLM layers)", "FFFFFF",
         "All 3 LLMs see\nredacted text only"),
        ("15803D", "6. RESTORE\n(Output handler)", "FFFFFF",
         "Leak detect then\nde-redact per chunk"),
        ("15803D", "7. AUDIT\n(DynamoDB)", "FFFFFF",
         "Entity types & counts\nNever values"),
    ]

    t = doc.add_table(rows=2, cols=len(stages))
    t.style = "Table Grid"

    for col_idx, (fill, title, tc, desc) in enumerate(stages):
        top = t.cell(0, col_idx)
        bot = t.cell(1, col_idx)
        _shade_cell(top, fill)
        _shade_cell(bot, "F0FDF4")
        rgb_tc = RGBColor(int(tc[0:2], 16), int(tc[2:4], 16), int(tc[4:6], 16))
        _cell_para(top, title, bold=True, size=8, color=rgb_tc, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(bot, desc, size=7, color=TEXT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)


# ---------------------------------------------------------------------------
# Diagram: Authentication Flow
# ---------------------------------------------------------------------------

def diagram_auth(doc: Document) -> None:
    steps = [
        ("0891B2", "Client authenticates", "Cognito User Pool\nReturns RS256 JWT (60 min TTL)\n+ Refresh token (30 days, never sent to API)"),
        ("0891B2", "API Request", "POST /chat\nAuthorization: Bearer <JWT>"),
        ("0891B2", "Lambda Authoriser", "Validates JWT signature vs Cognito JWKS\nChecks: exp, iss, aud, token_use\nInvalid  ->  401 at edge (no SQS, no pipeline cost)"),
        ("0891B2", "Validated", "user_sub (Cognito UUID) injected into context\nFlows through PipelineEnvelope\nNever: email, phone, name, or any PII claim"),
        ("0891B2", "Pipeline uses user_sub for", "Bouncer: banned-user check, rate-limit keys\nClassifier: session history lookup (last 3 turns)\nStrategist: tier-based vendor selection\nAudit: per-request attribution"),
    ]

    t = doc.add_table(rows=len(steps), cols=2)
    t.style = "Table Grid"

    for r_idx, (fill, title, detail) in enumerate(steps):
        c0, c1 = t.rows[r_idx].cells[0], t.rows[r_idx].cells[1]
        _shade_cell(c0, fill if r_idx % 2 == 0 else "0369A1")
        _shade_cell(c1, SH_AUTH if r_idx % 2 == 0 else "E0F2FE")
        _cell_para(c0, title, bold=True, size=9, color=WHITE)
        _cell_para(c1, detail, size=8, color=TEXT_DARK)


# ---------------------------------------------------------------------------
# Diagram: Compute Topology (ECS vs Lambda)
# ---------------------------------------------------------------------------

COMPUTE_TABLE = [
    ("Orchestrator",        "ECS Fargate",  "Persistent Redis + Presidio connections; VPC required"),
    ("Presidio Sidecar",    "ECS Fargate",  "200MB spaCy model must stay resident; 1-2ms VPC hop"),
    ("Bouncer (Haiku gate)","ECS Fargate",  "200ms latency budget; can't absorb Lambda cold start"),
    ("Intent Classifier",   "ECS Fargate",  "Persistent MCP clients; in-memory taxonomy cache"),
    ("Vendor Adapters",     "ECS Fargate",  "Streaming HTTP responses need persistent connection"),
    ("Admin Dashboard API", "ECS Fargate",  "Persistent Redis + DynamoDB; read-heavy, low concurrency"),
    ("WebSocket Server",    "ECS Fargate",  "Lambda cannot hold persistent long-lived connections"),
    ("Rule Gate",           "In-process",   "Pure Python, microseconds, no LLM — runs inside Orchestrator"),
    ("SQS Consumers",       "Lambda",       "Native SQS trigger; scales with queue depth"),
    ("JWT Authoriser",      "Lambda",       "Stateless JWT validation; millisecond execution"),
    ("Routing Strategist",  "Lambda*",      "80-90% pure rule lookups; Phase 2 migration from ECS"),
]

def diagram_compute(doc: Document) -> None:
    t = doc.add_table(rows=len(COMPUTE_TABLE) + 1, cols=3)
    t.style = "Table Grid"

    for i, label in enumerate(["Component", "Platform", "Rationale"]):
        c = t.cell(0, i)
        _shade_cell(c, C_HEADER)
        _cell_para(c, label, bold=True, size=9, color=WHITE)

    for r_idx, (comp, platform, rationale) in enumerate(COMPUTE_TABLE, start=1):
        row = t.rows[r_idx]
        fill = SH_INFRA if r_idx % 2 == 0 else "FFFFFF"
        plat_fill = "D1FAE5" if "Fargate" in platform else ("DBEAFE" if "Lambda" in platform else "FEF3C7")
        _shade_cell(row.cells[0], fill)
        _shade_cell(row.cells[1], plat_fill)
        _shade_cell(row.cells[2], fill)
        _cell_para(row.cells[0], comp, bold=True, size=9, color=TEXT_DARK)
        _cell_para(row.cells[1], platform, bold=True, size=9, color=TEXT_DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[2], rationale, size=8, color=TEXT_DARK)


# ---------------------------------------------------------------------------
# Diagram: VPC Topology
# ---------------------------------------------------------------------------

def diagram_vpc(doc: Document) -> None:
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    outer = t.cell(0, 0)
    _shade_cell(outer, "EFF6FF")

    outer.paragraphs[0].paragraph_format.space_before = Pt(4)
    _cell_para(outer, "AWS VPC  —  ap-southeast-1  (Singapore)", bold=True, size=10,
               color=RGBColor(0x1D, 0x4E, 0xD8))

    groups = [
        ("Public Subnet", "0891B2", [
            "API Gateway  (managed, VPC endpoint)",
            "Application Load Balancer  (Admin dashboard)",
            "CloudFront distribution  (React SPA)",
        ]),
        ("Private Subnet A  (application)", "0E3B2A", [
            "ECS Fargate: Orchestrator service",
            "ECS Fargate: Bouncer / Haiku gate service",
            "ECS Fargate: Intent Classifier service",
            "ECS Fargate: Routing Strategist service",
            "ECS Fargate: Vendor Adapter service",
            "ECS Fargate: Presidio Sidecar  (port 8080)",
            "ECS Fargate: Admin Dashboard API  (port 8000)",
            "ECS Fargate: WebSocket Server",
            "Lambda: JWT Authoriser  (VPC-attached)",
            "Lambda: SQS Consumers",
        ]),
        ("Private Subnet B  (data)", "7C3AED", [
            "ElastiCache Redis  (session cache, token vault, rate limits)",
            "SQS: Incoming queue  /  Escalation queue  /  Dead-letter queue",
            "DynamoDB: Routing rules, audit log, review log, intent taxonomy",
        ]),
        ("AWS Managed Services  (via VPC endpoints)", "475569", [
            "Amazon Bedrock  (bedrock-runtime endpoint)  —  all LLM calls",
            "Amazon Cognito  (User Pool in ap-southeast-1)",
            "CloudWatch Metrics + Logs  |  X-Ray  |  S3 archival",
            "AWS Cloud Map  (service discovery: presidio.internal)",
        ]),
    ]

    for group_name, fill, items in groups:
        gp = outer.add_paragraph()
        gp.paragraph_format.space_before = Pt(6)
        gp.paragraph_format.space_after = Pt(1)
        gr = gp.add_run(group_name)
        gr.bold = True
        gr.font.size = Pt(9)
        gr.font.color.rgb = RGBColor(
            int(fill[0:2], 16), int(fill[2:4], 16), int(fill[4:6], 16)
        )
        for item in items:
            ip = outer.add_paragraph()
            ip.paragraph_format.space_before = Pt(0)
            ip.paragraph_format.space_after = Pt(0)
            ir = ip.add_run("    " + item)
            ir.font.size = Pt(8)
            ir.font.color.rgb = TEXT_DARK

    ep = outer.add_paragraph()
    ep.paragraph_format.space_before = Pt(4)


# ---------------------------------------------------------------------------
# Diagram: PipelineEnvelope data contract
# ---------------------------------------------------------------------------

ENVELOPE_FIELDS = [
    ("correlation_id",        "str",   "UUID set at Orchestrator entry; threads all logs, spans, audit records"),
    ("user_sub",              "str",   "Cognito UUID from validated JWT. Never email/name/phone."),
    ("session_id",            "str",   "Identifies the conversation for session-history lookup"),
    ("redacted_message",      "str",   "Message after Presidio redaction. Only form that leaves the Orchestrator."),
    ("raw_message_hash",      "str",   "SHA-256 of original input for audit. Raw text never propagated."),
    ("entity_types_redacted", "list",  "e.g. ['SG_NRIC', 'EMAIL_ADDRESS']. Types only, never values."),
    ("entity_count",          "int",   "Count of redacted entities. 0 when no PII detected."),
    ("was_redacted",          "bool",  "True if any entity was found and tokenised."),
    ("timestamp",             "str",   "ISO 8601 UTC time when message entered the pipeline."),
    ("bedrock_region",        "str",   "ap-southeast-1. Baked into envelope; verified by policy engine."),
    ("source_ip",             "str",   "Client IP from API Gateway for audit. Not propagated to LLMs."),
]

def diagram_envelope(doc: Document) -> None:
    t = doc.add_table(rows=len(ENVELOPE_FIELDS) + 1, cols=3)
    t.style = "Table Grid"

    for i, label in enumerate(["Field", "Type", "Purpose"]):
        c = t.cell(0, i)
        _shade_cell(c, C_HEADER)
        _cell_para(c, label, bold=True, size=9, color=WHITE)

    for r_idx, (field, ftype, purpose) in enumerate(ENVELOPE_FIELDS, start=1):
        row = t.rows[r_idx]
        fill = SH_ORCH if r_idx % 2 == 0 else "FFFFFF"
        _shade_cell(row.cells[0], fill)
        _shade_cell(row.cells[1], fill)
        _shade_cell(row.cells[2], fill)
        _cell_para(row.cells[0], field, bold=True, size=9,
                   color=RGBColor(0x0E, 0x3B, 0x2A))
        _cell_para(row.cells[1], ftype, size=8, color=BRAND_MID,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[2], purpose, size=8, color=TEXT_DARK)


# ---------------------------------------------------------------------------
# Diagram: ADR summary table
# ---------------------------------------------------------------------------

ADRS = [
    ("ADR-001", "All LLM calls via Bedrock",
     "Single IAM-enforced data residency, single billing surface, no vendor API keys. "
     "No direct Anthropic/OpenAI imports. CI enforces with forbidden-imports check."),
    ("ADR-002", "Presidio as persistent ECS sidecar",
     "200MB spaCy model must stay resident. Lambda cold-start (3-8s) is unacceptable. "
     "ECS keeps model warm; 1-2ms VPC-internal HTTP hop. Independent scaling & deployment."),
    ("ADR-003", "Cognito JWT validated at API Gateway edge",
     "Unauthenticated requests rejected before SQS and any pipeline cost. "
     "user_sub (UUID) is the only identity propagated — no PII-bearing claims downstream."),
    ("ADR-004", "Escalation split: Review Log vs Graceful Degradation Ladder",
     "Single queue conflated real-time response and async audit. "
     "Now: Review Log (async DynamoDB) + Ladder (Hedge -> Clarify -> Fallback -> Hand-off)."),
]

def diagram_adrs(doc: Document) -> None:
    t = doc.add_table(rows=len(ADRS) + 1, cols=3)
    t.style = "Table Grid"

    for i, label in enumerate(["ADR", "Decision", "Key Rationale"]):
        c = t.cell(0, i)
        _shade_cell(c, C_HEADER)
        _cell_para(c, label, bold=True, size=9, color=WHITE)

    for r_idx, (adr_id, decision, rationale) in enumerate(ADRS, start=1):
        row = t.rows[r_idx]
        fill = SH_ALT if r_idx % 2 == 0 else "FFFFFF"
        for c in row.cells:
            _shade_cell(c, fill)
        _cell_para(row.cells[0], adr_id, bold=True, size=9, color=RGBColor(0x0E, 0x3B, 0x2A),
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[1], decision, bold=True, size=9, color=TEXT_DARK)
        _cell_para(row.cells[2], rationale, size=8, color=TEXT_DARK)


# ---------------------------------------------------------------------------
# Tech stack table
# ---------------------------------------------------------------------------

STACK = [
    ("Language & Runtime",   "Python 3.12  |  FastAPI  |  FastMCP"),
    ("LLM API",              "AWS Bedrock  —  all LLM calls (no direct Anthropic/OpenAI clients)"),
    ("Models",               "Claude Haiku 4.5  (Bouncer, Strategist arbitration)\n"
                             "Claude Sonnet 4.6  (Classifier deep path, default vendor)\n"
                             "Cohere Embed Multilingual v3  (Classifier fast path)"),
    ("Model IDs",            "global.anthropic.claude-haiku-4-5-20251001-v1:0  (cross-region)\n"
                             "global.anthropic.claude-sonnet-4-6  (cross-region)\n"
                             "cohere.embed-multilingual-v3"),
    ("Message Queue",        "AWS SQS  —  Incoming, Escalation, Dead-letter queues"),
    ("Cache / Vault",        "ElastiCache Redis  —  session history, token vault, rate limits"),
    ("Config Store",         "DynamoDB  —  routing rules, audit log, review log, intent taxonomy"),
    ("PII Detection",        "Microsoft Presidio  +  custom SG recognisers (NRIC, FIN, UEN)"),
    ("Vendor Adapter",       "LiteLLM (pinned version, configured against Bedrock)"),
    ("Compute",              "ECS Fargate  (latency-sensitive, stateful)\n"
                             "AWS Lambda  (stateless, bursty)"),
    ("Auth",                 "Amazon Cognito User Pool  +  JWT RS256\n"
                             "Lambda Authoriser at API Gateway edge"),
    ("Networking",           "VPC  ap-southeast-1  |  Bedrock VPC endpoint  |  Cloud Map"),
    ("Region",               "ap-southeast-1  (Singapore)  —  single region by design"),
    ("IaC",                  "Terraform  (all infra under infra/terraform/)"),
    ("Observability",        "CloudWatch Metrics + Logs  |  X-Ray distributed tracing  |  S3 archival"),
    ("Frontend",             "React SPA  |  S3 + CloudFront  |  Cognito auth"),
    ("Dependency Mgr",       "uv  (uv.lock committed)  —  no requirements.txt"),
    ("Data Validation",      "Pydantic v2  (all handoff models)"),
    ("Test Framework",       "pytest + pytest-asyncio  |  LocalStack  |  fakeredis"),
]

def diagram_stack(doc: Document) -> None:
    t = doc.add_table(rows=len(STACK) + 1, cols=2)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, label in enumerate(["Category", "Technology / Detail"]):
        _shade_cell(hdr[i], C_HEADER)
        _cell_para(hdr[i], label, bold=True, size=9, color=WHITE)

    for r_idx, (cat, detail) in enumerate(STACK, start=1):
        row = t.rows[r_idx]
        fill = SH_ALT if r_idx % 2 == 0 else "FFFFFF"
        _shade_cell(row.cells[0], fill)
        _shade_cell(row.cells[1], fill)
        _cell_para(row.cells[0], cat, bold=True, size=9, color=TEXT_DARK)
        _cell_para(row.cells[1], detail, size=9, color=TEXT_DARK)


# ---------------------------------------------------------------------------
# Diagram: MCP Usage
# ---------------------------------------------------------------------------

MCP_TABLE = [
    ("Inbound  (Router as MCP server)",
     "FastMCP on ECS Fargate",
     "route_message, get_routing_status",
     "External agents can submit messages via the MCP protocol"),
    ("Classifier  ->  taxonomy server",
     "DynamoDB-backed MCP server",
     "get_intent_taxonomy",
     "Full intent hierarchy fetched per classification request"),
    ("Classifier  ->  session server",
     "Redis-backed MCP server",
     "get_session_history",
     "Last 3 turns for pronoun resolution in resolved_message"),
    ("Classifier  ->  guardrails server",
     "Bedrock Guardrails MCP server",
     "check_guardrails",
     "Defence-in-depth content check on classifier input"),
    ("Classifier  ->  entity context server",
     "Redis-backed MCP server",
     "get_entity_context",
     "Known entities for this user (for disambiguation)"),
    ("Strategist  ->  vendor health server",
     "CloudWatch-backed MCP server",
     "check_vendor_health",
     "Issued in parallel with rule lookup on every request"),
    ("Strategist  ->  routing rules server",
     "DynamoDB-backed MCP server",
     "get_routing_rules",
     "Current intent->vendor mapping with TTL cache"),
    ("Bouncer",
     "—",
     "No MCP tools",
     "Rule gate + Haiku only. No tool calls in the 200ms budget."),
]

def diagram_mcp(doc: Document) -> None:
    t = doc.add_table(rows=len(MCP_TABLE) + 1, cols=4)
    t.style = "Table Grid"
    hdrs = ["Layer / Direction", "Server Type", "Tool(s)", "Purpose"]
    for i, label in enumerate(hdrs):
        c = t.cell(0, i)
        _shade_cell(c, C_HEADER)
        _cell_para(c, label, bold=True, size=9, color=WHITE)

    for r_idx, (layer, server, tools, purpose) in enumerate(MCP_TABLE, start=1):
        row = t.rows[r_idx]
        fill = SH_ALT if r_idx % 2 == 0 else "FFFFFF"
        for c in row.cells:
            _shade_cell(c, fill)
        _cell_para(row.cells[0], layer, bold=True, size=8, color=TEXT_DARK)
        _cell_para(row.cells[1], server, size=8, color=TEXT_DARK)
        _cell_para(row.cells[2], tools, size=8,
                   color=RGBColor(0x7C, 0x3A, 0xED) if tools != "No MCP tools" else BRAND_MID)
        _cell_para(row.cells[3], purpose, size=8, color=TEXT_DARK)


# ---------------------------------------------------------------------------
# Diagram: Non-negotiables summary
# ---------------------------------------------------------------------------

NON_NEG = [
    ("1", "All LLM calls via Bedrock",
     "No import anthropic / import openai. CI forbidden-imports check enforces this."),
    ("2", "JWT validated at edge",
     "API Gateway authoriser rejects unauthenticated requests. Pipeline never re-validates."),
    ("3", "Single redaction policy",
     "One Presidio pass at orchestrator entry. Applied uniformly to Haiku, Sonnet, and vendor Claude."),
    ("4", "Fail-open Bouncer",
     "200ms timeout -> allowed=True, timed_out=True. Never reject on timeout. Tests enforce this."),
    ("5", "Correlation ID threads everything",
     "Set as ContextVar at Orchestrator entry. Never passed as a function argument."),
    ("6", "No PII in logs",
     "safe_log() with allowlist. Errors logged as type names, never error.message."),
    ("7", "Bedrock invocation logging disabled",
     "loggingConfig: null on every Bedrock invocation. Application audit is the trail."),
    ("8", "Escalate, don't guess",
     "Confidence below threshold -> human review queue. Never pick a default vendor blindly."),
    ("9", "Data residency at IAM layer",
     "apac.* / global.* inference profile IDs only. IAM condition locks region. No raw model IDs."),
    ("10", "Raw message never leaves Orchestrator",
     "Only redacted_message in PipelineEnvelope passes downstream. raw_message_hash for audit only."),
]

def diagram_non_neg(doc: Document) -> None:
    t = doc.add_table(rows=len(NON_NEG) + 1, cols=3)
    t.style = "Table Grid"
    for i, label in enumerate(["#", "Rule", "How it is enforced"]):
        c = t.cell(0, i)
        _shade_cell(c, C_HEADER)
        _cell_para(c, label, bold=True, size=9, color=WHITE)

    for r_idx, (num, rule, enforcement) in enumerate(NON_NEG, start=1):
        row = t.rows[r_idx]
        fill = SH_ALT if r_idx % 2 == 0 else "FFFFFF"
        for c in row.cells:
            _shade_cell(c, fill)
        _cell_para(row.cells[0], num, bold=True, size=10, color=RGBColor(0x0E, 0x3B, 0x2A),
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[1], rule, bold=True, size=9, color=TEXT_DARK)
        _cell_para(row.cells[2], enforcement, size=8, color=TEXT_DARK)


# ---------------------------------------------------------------------------
# Main build
# ---------------------------------------------------------------------------

def build() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    _set_styles(doc)

    # ------------------------------------------------------------------
    # Cover
    # ------------------------------------------------------------------
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(72)
    cp.paragraph_format.space_after  = Pt(6)
    r = cp.add_run("Evidor.ai  ·  AI Router")
    r.font.name = "Calibri"; r.font.size = Pt(13)
    r.font.color.rgb = BRAND_MID; r.bold = True

    ch = doc.add_paragraph()
    ch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ch.paragraph_format.space_after = Pt(4)
    rh = ch.add_run("Technical Architecture\nDocument")
    rh.font.name = "Calibri"; rh.font.size = Pt(28)
    rh.bold = True; rh.font.color.rgb = BRAND_DARK

    cs = doc.add_paragraph()
    cs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cs.paragraph_format.space_after = Pt(72)
    rs = cs.add_run("Version 1.0  ·  08 June 2026  ·  Confidential & Internal")
    rs.font.name = "Calibri"; rs.font.size = Pt(10)
    rs.font.color.rgb = BRAND_MID

    doc.add_page_break()

    # ------------------------------------------------------------------
    # TOC
    # ------------------------------------------------------------------
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # ==================================================================
    # 1. Executive Summary
    # ==================================================================
    H(doc, "1  Executive Summary", 1)
    P(doc,
      "The AI Router is a multi-layer message processing pipeline that intercepts incoming chat "
      "messages, validates and sanitises them, classifies intent, selects an appropriate AI vendor, "
      "and streams a response back to the user — all while enforcing strict data residency, PII "
      "protection, and content safety requirements for the Singapore market.")

    P(doc,
      "Every message passes through five sequential processing layers: Bouncer (content safety "
      "gate), Intent Classifier (message understanding), Routing Strategist (vendor selection), "
      "Vendor Adapter (Bedrock LLM invocation), and an Output Handler that restores redacted "
      "entities and audits the interaction. An Orchestrator drives the pipeline end-to-end.")

    note(doc, "Key design principle:",
         "All LLM calls go through Amazon Bedrock. No direct Anthropic or OpenAI API clients "
         "are permitted. PII is redacted once at pipeline entry and never re-introduced downstream. "
         "The raw message never leaves the Orchestrator process.", "0E3B2A")

    # ==================================================================
    # 2. Technology Stack
    # ==================================================================
    H(doc, "2  Technology Stack", 1)
    P(doc, "The complete set of technologies used across all layers of the AI Router.")
    diagram_stack(doc)
    doc.add_paragraph()

    # ==================================================================
    # 3. High-Level Architecture
    # ==================================================================
    H(doc, "3  High-Level Architecture", 1)
    P(doc,
      "The diagram below shows the complete request flow from client to response. Each coloured "
      "block is a distinct processing layer. The pipeline is sequential — each layer receives a "
      "structured result object from the previous layer and returns its own result to the Orchestrator.")

    note(doc, "Correlation ID:",
         "A UUID is assigned at the very first step (Orchestrator entry) and threads through "
         "every log line, queue message, database write, and span in the entire pipeline. "
         "It is set as a Python ContextVar — never passed as a function argument.", "0E3B2A")

    diagram_pipeline_overview(doc)
    diagram_caption(doc, "Figure 1  —  End-to-end request flow through the AI Router pipeline")

    # ==================================================================
    # 4. Authentication & Security Gate
    # ==================================================================
    H(doc, "4  Authentication & Security Gate", 1)
    P(doc,
      "Authentication is handled entirely at the API Gateway edge using Amazon Cognito and JWT "
      "access tokens. Unauthenticated requests are rejected before any pipeline code runs, "
      "before any message reaches SQS, and before any compute or Bedrock cost is incurred.")

    H(doc, "4.1  Authentication Flow", 2)
    diagram_auth(doc)
    diagram_caption(doc, "Figure 2  —  Cognito JWT authentication flow and user_sub propagation")
    doc.add_paragraph()

    H(doc, "4.2  Token Lifetimes", 2)
    B(doc, "Access token", "RS256 JWT, 60-minute TTL. Sent with every API request.")
    B(doc, "Refresh token", "30-day TTL. Stored client-side only. Never sent to the API.")
    B(doc, "user_sub", "Cognito UUID. The only identity field propagated into the pipeline. "
           "No email, phone, or name ever enters downstream components.")

    H(doc, "4.3  Why a Lambda Authoriser over the built-in Cognito authoriser", 2)
    B(doc, "Banned-user check at edge",
      "The Lambda authoriser can reject a banned user_sub before SQS, eliminating cost for every subsequent request.")
    B(doc, "Custom context injection",
      "User tier and region are injected into request context once, avoiding extra DB lookups in downstream layers.")
    B(doc, "Stateless validation",
      "JWT validation against Cognito JWKS is pure computation — millisecond execution, no shared state.")

    # ==================================================================
    # 5. PII Redaction & Data Privacy
    # ==================================================================
    H(doc, "5  PII Redaction & Data Privacy", 1)
    P(doc,
      "PII is detected and tokenised before any LLM sees the message. A single redaction pass at "
      "Orchestrator entry applies uniformly to every downstream LLM call — Haiku, Sonnet, and the "
      "vendor adapter. There is no per-layer redaction.")

    H(doc, "5.1  Redaction Lifecycle", 2)
    diagram_redaction(doc)
    diagram_caption(doc, "Figure 3  —  PII redaction lifecycle from ingest to audit")
    doc.add_paragraph()

    H(doc, "5.2  What Gets Redacted", 2)
    P(doc, "High-sensitivity entities — always redacted:", italic=True)
    B(doc, "SG_NRIC, SG_FIN, SG_UEN, SG_PASSPORT", "Singapore national identifiers (custom Presidio recognisers)")
    B(doc, "CREDIT_CARD, IBAN_CODE", "Financial instrument numbers")

    P(doc, "Contextual entities — NOT redacted (needed for classification accuracy):", italic=True)
    B(doc, "LOCATION, PERSON, DATE_TIME", "Needed by the Classifier to understand query context")
    B(doc, "PHONE_NUMBER, EMAIL_ADDRESS", "Left as-is for transactional queries (password reset, callbacks)")

    H(doc, "5.3  Token Vault", 2)
    B(doc, "Key format", "vault:{correlation_id}:{token}  —  scoped per request")
    B(doc, "TTL", "5 minutes — security boundary, not a data-retention window")
    B(doc, "Eager cleanup", "Vault keys are deleted in the Orchestrator finally block after the response is sent")
    B(doc, "Streaming restoration", "200-char buffer with 50-char safety margin for mid-entity chunk splits")

    H(doc, "5.4  Output Leak Detector", 2)
    P(doc,
      "After the vendor response is restored, the output handler scans each chunk for PII entity "
      "types that were NOT present in the original input. Any unexpected PII (hallucinated entity "
      "values) is stripped before delivery to the user.")

    warn(doc, "Bedrock invocation logging is disabled.",
         "Terraform sets loggingConfig: null on every Bedrock invocation to prevent PII "
         "from leaking into AWS-managed logs. Application-level structlog is the audit trail.")

    # ==================================================================
    # 6. Processing Layers in Detail
    # ==================================================================
    H(doc, "6  Processing Layers in Detail", 1)

    # --- 6.1 Bouncer ---
    H(doc, "6.1  Bouncer  (Layer 1)", 2)
    P(doc,
      "The Bouncer is the first processing gate after the Orchestrator constructs the "
      "PipelineEnvelope. It has a hard 200ms total budget shared between two sub-stages. "
      "Its defining characteristic is fail-open on timeout — the pipeline always continues.")

    H(doc, "Bouncer Decision Logic", 3)
    diagram_bouncer(doc)
    diagram_caption(doc, "Figure 4  —  Bouncer two-stage gate with all decision paths")
    doc.add_paragraph()

    danger(doc, "Critical: Fail-open is not optional.",
           "When the 200ms budget is exceeded, the Bouncer must return allowed=True with "
           "timed_out=True. Raising an exception or returning allowed=False on timeout is a "
           "spec violation. The test suite enforces this explicitly.")

    B(doc, "Rule Gate (Stage 1)", "Runs in microseconds. Length check, prompt injection regex, "
      "banned user_sub lookup in Redis, rate limit check. If the rule gate rejects, Haiku is never called.")
    B(doc, "Haiku micro-classifier (Stage 2)", "Max 50 tokens output. Returns {pass, reason, confidence}. "
      "Uses the global Haiku 4.5 inference profile.")
    B(doc, "Escalation vs Rejection", "Haiku can only escalate (allowed=True, escalate=True). "
      "Only the rule gate can hard-reject (allowed=False). Haiku errors always fail-open.")

    # --- 6.2 Classifier ---
    H(doc, "6.2  Intent Classifier  (Layer 2)", 2)
    P(doc,
      "The Classifier determines what the user wants. It uses a two-path design: a fast "
      "embedding-similarity path for clear-cut messages, and a deep Sonnet path with MCP tools "
      "for nuanced or multi-domain queries.")

    H(doc, "Two-Path Design", 3)
    diagram_classifier(doc)
    diagram_caption(doc, "Figure 5  —  Classifier fast path (embedding) vs deep path (Sonnet + MCP)")
    doc.add_paragraph()

    H(doc, "Intent Domains", 3)
    B(doc, "general_qa", "General knowledge questions -> Claude Sonnet 4.6")
    B(doc, "code_assistance", "Coding, debugging, programming help -> Claude Sonnet 4.6")
    B(doc, "simple_transactional", "Quick lookups, prices, hours -> Claude Haiku (cheaper, faster)")
    B(doc, "out_of_scope", "Action requests the router cannot fulfil -> escalate")
    B(doc, "ambiguous", "Cannot determine intent with sufficient confidence -> escalate")

    note(doc, "Resolved message:",
         "The Classifier produces a resolved_message with pronouns disambiguated using session "
         "history (last 3 turns). This is still the redacted form — never the raw message.", "7C3AED")

    # --- 6.3 Strategist ---
    H(doc, "6.3  Routing Strategist  (Layer 3)", 2)
    P(doc,
      "The Strategist selects the AI vendor and constructs the RoutingPlan. It operates on "
      "three confidence tiers and always runs the policy engine after vendor selection to enforce "
      "data residency and compliance rules.")

    H(doc, "Confidence Tier Routing", 3)
    diagram_strategist(doc)
    diagram_caption(doc, "Figure 6  —  Routing Strategist confidence tiers and actions")
    doc.add_paragraph()

    H(doc, "Policy Engine", 3)
    B(doc, "Data residency", "SG users must be routed to Bedrock ap-southeast-1 only. "
      "Enforced at both IAM level and policy engine. Non-compliant selections are blocked and escalated.")
    B(doc, "MAS-regulated queries", "Financial advice intents trigger compliance flags. "
      "Certain sub-intents are blocked entirely or require special vendor configuration.")
    B(doc, "Cross-jurisdiction flags", "Overseas content for SG users is flagged and may trigger hand-off.")
    B(doc, "Fallback chain", "Per-intent vendor timeouts (3s Haiku, 8s Sonnet). "
      "Each step has its own retry count and exponential backoff.")

    # --- 6.4 Vendor Adapters ---
    H(doc, "6.4  Vendor Adapters  (Layer 4)", 2)
    P(doc,
      "All vendors are unified behind Amazon Bedrock. LiteLLM provides a consistent interface "
      "across Claude, Llama, and Mistral models. Streaming is enabled by default for chat-like UX.")

    B(doc, "LiteLLM adapter", "Handles multi-model API differences, retries, and cost tracking. "
      "Configured against Bedrock inference profiles only.")
    B(doc, "Streaming", "InvokeModelWithResponseStream. Chunks are sent to the Orchestrator "
      "for leak detection and vault restoration before forwarding to the WebSocket.")
    B(doc, "Model IDs", "Cross-region inference profiles only (global.* or apac.*). "
      "Raw anthropic.* model IDs fail with 'on-demand throughput not supported'.")

    warn(doc, "LiteLLM pinning:",
         "Pin to a known-good version in pyproject.toml. Versions 1.82.7 and 1.82.8 had "
         "security advisories. Do not float the version.")

    # --- 6.5 Orchestrator ---
    H(doc, "6.5  Orchestrator", 2)
    P(doc,
      "The Orchestrator is the pipeline driver. It dequeues messages from SQS, drives each "
      "layer in sequence, and ensures audit and cleanup happen in a finally block regardless "
      "of success or failure.")

    B(doc, "correlation_id", "Set as a ContextVar at the very first step. Every log line, "
      "span, and audit record inherits it automatically — never passed as a function argument.")
    B(doc, "PipelineEnvelope", "Built once after redaction. Immutable. Contains the "
      "redacted_message and raw_message_hash. Never modified by downstream layers.")
    B(doc, "Finally block", "Runs on both success and failure paths. Writes DynamoDB audit "
      "record and deletes Redis vault keys. Emits CloudWatch metrics.")
    B(doc, "Fire-and-forget metrics", "emit_orchestrator(), emit_classifier(), emit_strategist() "
      "are scheduled via asyncio.create_task() in the finally block.")

    # ==================================================================
    # 7. Data Contracts
    # ==================================================================
    H(doc, "7  Data Contracts", 1)
    P(doc,
      "Every layer communicates through strongly-typed Pydantic v2 models. The PipelineEnvelope "
      "is the central carrier — built once after redaction and passed immutably through every "
      "downstream layer.")

    H(doc, "7.1  PipelineEnvelope", 2)
    diagram_envelope(doc)
    diagram_caption(doc, "Figure 7  —  PipelineEnvelope fields: the immutable request carrier")
    doc.add_paragraph()

    H(doc, "7.2  Layer Output Contracts", 2)
    contracts = [
        ("BounceResult",     "Bouncer",     "allowed, reason, layer, confidence, escalate, timed_out"),
        ("ClassifiedIntent", "Classifier",  "intent, sub_intent, domain, confidence, entities, resolved_message, multi_domain, escalate, reasoning"),
        ("RoutingPlan",      "Strategist",  "primary_vendor, fallback_chain, context, applied_policies, policy_modified, blocked"),
        ("RedactionResult",  "Orchestrator","redacted_message, entity_types_found, entity_count, was_redacted, correlation_id"),
    ]
    t = doc.add_table(rows=len(contracts) + 1, cols=3)
    t.style = "Table Grid"
    for i, label in enumerate(["Model", "Owner Layer", "Key Fields"]):
        c = t.cell(0, i)
        _shade_cell(c, C_HEADER)
        _cell_para(c, label, bold=True, size=9, color=WHITE)
    for r_idx, (model, owner, fields) in enumerate(contracts, start=1):
        row = t.rows[r_idx]
        fill = SH_ALT if r_idx % 2 == 0 else "FFFFFF"
        for c in row.cells:
            _shade_cell(c, fill)
        _cell_para(row.cells[0], model, bold=True, size=9, color=RGBColor(0x0E, 0x3B, 0x2A))
        _cell_para(row.cells[1], owner, size=9, color=TEXT_DARK)
        _cell_para(row.cells[2], fields, size=8, color=TEXT_DARK)
    doc.add_paragraph()

    # ==================================================================
    # 8. MCP (Model Context Protocol)
    # ==================================================================
    H(doc, "8  MCP  (Model Context Protocol)", 1)
    P(doc,
      "The router both exposes an MCP server (for external agents to submit messages) and "
      "consumes MCP servers as a client (for Classifier and Strategist tool calls). Each MCP "
      "server owns exactly one domain of data — taxonomy, session history, vendor health, etc.")

    note(doc, "Parallel tool calls:",
         "The Classifier issues all MCP tool calls via asyncio.gather() — never sequentially. "
         "Sequential calls would blow the classification latency budget.", "7C3AED")

    diagram_mcp(doc)
    diagram_caption(doc, "Figure 8  —  MCP server usage across all pipeline layers")
    doc.add_paragraph()

    # ==================================================================
    # 9. Infrastructure & Compute
    # ==================================================================
    H(doc, "9  Infrastructure & Compute", 1)
    P(doc,
      "The compute model follows one principle: ECS Fargate for latency-sensitive or stateful "
      "components, Lambda for stateless bursty workloads. The Presidio sidecar's 200MB spaCy "
      "model is the key driver of the ECS-heavy topology.")

    H(doc, "9.1  Compute Decisions per Component", 2)
    diagram_compute(doc)
    diagram_caption(doc, "Figure 9  —  Component compute platform and rationale")
    doc.add_paragraph()

    H(doc, "9.2  VPC Topology", 2)
    P(doc,
      "All pipeline components run inside a single VPC in ap-southeast-1. Bedrock is accessed "
      "via a VPC endpoint — no LLM traffic leaves the VPC over the public internet.")
    diagram_vpc(doc)
    diagram_caption(doc, "Figure 10  —  VPC topology: public, private application, and private data subnets")
    doc.add_paragraph()

    H(doc, "9.3  Phased Rollout", 2)
    B(doc, "Phase 1 (current)",
      "All layers on ECS Fargate. Single service per logical layer. Simplest to debug and observe.")
    B(doc, "Phase 2",
      "Migrate stateless layers (Rule Gate, SQS Consumers, Routing Strategist) to Lambda once "
      "traffic data justifies it. ECS layers unchanged.")
    B(doc, "Phase 3",
      "Add Lambda provisioned concurrency where cold starts appear in CloudWatch p99 metrics. "
      "Add per-layer SQS queues for horizontal scaling under load.")

    # ==================================================================
    # 10. Observability
    # ==================================================================
    H(doc, "10  Observability", 1)

    H(doc, "10.1  Structured Logging", 2)
    P(doc,
      "All logging uses structlog with an allowlist-based context processor. Only fields in the "
      "allowlist are emitted. The correlation_id is auto-injected from the ContextVar into every "
      "log line — it is never passed as an explicit argument.")
    note(doc, "No PII in logs.",
         "Use safe_log() from shared/logging.py. Errors are logged as type(e).__name__, never "
         "as str(e) which may contain echoed user input. This is non-negotiable.", "DC2626")

    H(doc, "10.2  CloudWatch Metrics", 2)
    metrics = [
        ("AIRouter/Orchestrator", "RequestCount, RequestLatencyMs, ErrorCount, EscalationCount"),
        ("AIRouter/Bouncer",      "BouncerTimeout, BouncerError, PassCount, EscalationCount, AvgConfidence"),
        ("AIRouter/Classifier",   "FastPathCount, DeepPathCount, EscalationCount, AvgConfidence"),
        ("AIRouter/Strategist",   "DeterministicCount, ArbitrationCount, PolicyBlocked, FallbackUsed"),
        ("AIRouter/Admin",        "AdminRequestCount, AdminErrorCount, EscalationQueueDepth"),
    ]
    t = doc.add_table(rows=len(metrics) + 1, cols=2)
    t.style = "Table Grid"
    for i, label in enumerate(["Namespace", "Key Metrics"]):
        c = t.cell(0, i)
        _shade_cell(c, C_HEADER)
        _cell_para(c, label, bold=True, size=9, color=WHITE)
    for r_idx, (ns, mlist) in enumerate(metrics, start=1):
        row = t.rows[r_idx]
        fill = SH_ALT if r_idx % 2 == 0 else "FFFFFF"
        _shade_cell(row.cells[0], fill); _shade_cell(row.cells[1], fill)
        _cell_para(row.cells[0], ns, bold=True, size=9, color=TEXT_DARK)
        _cell_para(row.cells[1], mlist, size=8, color=TEXT_DARK)
    doc.add_paragraph()

    H(doc, "10.3  Distributed Tracing", 2)
    B(doc, "AWS X-Ray", "Traces span API Gateway -> Orchestrator -> all layer calls. "
      "correlation_id is injected as a custom annotation on every segment.")
    B(doc, "Audit log", "DynamoDB audit table stores entity types, counts, vendor used, "
      "latency per layer, and policy flags per correlation_id. 365-day retention, S3-backed.")
    B(doc, "CloudTrail", "Captures every Bedrock InvokeModel call for compliance review.")

    # ==================================================================
    # 11. Architectural Decisions  (ADR Summaries)
    # ==================================================================
    H(doc, "11  Key Architectural Decisions", 1)
    P(doc,
      "The four formal Architecture Decision Records (ADRs) capture the most consequential "
      "design choices. Full rationale and alternatives considered are in docs/adr/.")

    diagram_adrs(doc)
    diagram_caption(doc, "Figure 11  —  ADR summary: four locked decisions")
    doc.add_paragraph()

    # ==================================================================
    # 12. Security Non-Negotiables
    # ==================================================================
    H(doc, "12  Security Non-Negotiables", 1)
    P(doc,
      "These rules are checked on every pull request by CI. Any violation fails the build. "
      "They cannot be suspended for debugging, performance, or convenience.")

    diagram_non_neg(doc)
    diagram_caption(doc, "Figure 12  —  Ten security rules enforced by CI and code review")
    doc.add_paragraph()

    # ==================================================================
    # 13. File Structure
    # ==================================================================
    H(doc, "13  File Structure", 1)
    P(doc, "Target directory layout for the entire codebase.")

    modules = [
        ("bouncer/",         "models.py, rule_gate.py, llm_classifier.py, bouncer.py, config.py"),
        ("classifier/",      "models.py, fast_path.py, deep_path.py, taxonomy.py, session_history.py, classifier.py"),
        ("strategist/",      "models.py, vendor_selector.py, policy_engine.py, fallback_chain.py, strategist.py"),
        ("redactor/",        "models.py, vault.py, recognisers.py, input_redactor.py, output_redactor.py, streaming_redactor.py, audit_logger.py"),
        ("adapters/",        "base.py, bedrock_client.py, litellm_adapter.py"),
        ("orchestrator/",    "orchestrator.py, sqs_consumer.py, envelope.py, presidio_client.py, vault.py, pipeline_driver.py, output_handler.py, audit.py, observability.py, websocket_server.py, config.py"),
        ("mcp_servers/",     "router_server.py, taxonomy_server.py, vendor_health_server.py, session_server.py, guardrails_server.py"),
        ("admin/",           "main.py, auth.py, models.py  +  routers/, services/"),
        ("presidio_sidecar/","main.py, Dockerfile, recognisers/"),
        ("shared/",          "models.py, logging.py, bedrock.py, correlation.py, errors.py, metrics.py"),
        ("infra/terraform/", "main.tf, iam.tf, ecs.tf, sqs.tf, dynamodb.tf, redis.tf, cognito.tf, cloudwatch.tf"),
        ("tests/",           "bouncer/, classifier/, strategist/, redactor/, orchestrator/  — mirroring source layout"),
    ]

    t = doc.add_table(rows=len(modules) + 1, cols=2)
    t.style = "Table Grid"
    for i, label in enumerate(["Directory", "Key Files"]):
        c = t.cell(0, i)
        _shade_cell(c, C_HEADER)
        _cell_para(c, label, bold=True, size=9, color=WHITE)
    for r_idx, (d, files) in enumerate(modules, start=1):
        row = t.rows[r_idx]
        fill = SH_ALT if r_idx % 2 == 0 else "FFFFFF"
        _shade_cell(row.cells[0], fill); _shade_cell(row.cells[1], fill)
        _cell_para(row.cells[0], d, bold=True, size=9, color=RGBColor(0x0E, 0x3B, 0x2A))
        _cell_para(row.cells[1], files, size=8, color=TEXT_DARK)

    # ------------------------------------------------------------------
    # Save
    # ------------------------------------------------------------------
    out = Path("reports/AI_Router_Technical_Architecture.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Saved -> {out}")


if __name__ == "__main__":
    build()
