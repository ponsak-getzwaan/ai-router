"""Generate the Evidor.ai Core Architecture Technical Report as a .docx file.

Run:
    uv run python scripts/generate_arch_report.py

Output:
    reports/architecture-report.docx

Edit freely in Word / Google Docs. To regenerate from scratch, just re-run.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Brand colours
# ---------------------------------------------------------------------------
FOREST = RGBColor(0x0E, 0x3B, 0x2A)
CREAM = RGBColor(0xFA, 0xFA, 0xF7)
AMBER = RGBColor(0xC7, 0x6E, 0x3A)
SLATE = RGBColor(0x1E, 0x2A, 0x3A)
MUTED = RGBColor(0x6B, 0x7C, 0x73)
LIGHT_GREEN_HEX = "E8F0EB"
LIGHT_AMBER_HEX = "F9EDE5"
FOREST_HEX = "0E3B2A"
CREAM_HEX = "FAFAF7"
BORDER_HEX = "D8DDD9"
ALT_ROW_HEX = "F0F4F2"

# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        cfg = kwargs.get(side, {})
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), cfg.get("val", "single"))
        el.set(qn("w:sz"), str(cfg.get("sz", 4)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), cfg.get("color", BORDER_HEX))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _para_shade(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _set_row_height(row, twips: int):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(twips))
    trPr.append(trHeight)


def _add_bookmark(para, name: str):
    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), "0")
    bm_start.set(qn("w:name"), name)
    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), "0")
    para._p.append(bm_start)
    para._p.append(bm_end)

# ---------------------------------------------------------------------------
# Document-level helpers
# ---------------------------------------------------------------------------

def _configure_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = SLATE

    for name, size, bold, color in [
        ("Heading 1", 18, True, FOREST),
        ("Heading 2", 13, True, FOREST),
        ("Heading 3", 11, True, SLATE),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(14 if size == 18 else 10)
        s.paragraph_format.space_after = Pt(4)

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def _add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def _add_h2(doc, text):
    return doc.add_heading(text, level=2)


def _add_h3(doc, text):
    return doc.add_heading(text, level=3)


def _add_para(doc, text, bold=False, italic=False, color=None, size=10, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _add_bullet(doc, text, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        rest = p.add_run(text)
        rest.font.name = "Calibri"
        rest.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
    return p


def _add_info_box(doc, text, bg_hex=LIGHT_GREEN_HEX, left_border_color=FOREST_HEX):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, bg_hex)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    _set_cell_border(cell,
                     left={"val": "single", "sz": 18, "color": left_border_color},
                     top={"val": "none", "sz": 0, "color": "auto"},
                     bottom={"val": "none", "sz": 0, "color": "auto"},
                     right={"val": "none", "sz": 0, "color": "auto"})
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def _add_amber_box(doc, text):
    return _add_info_box(doc, text, bg_hex=LIGHT_AMBER_HEX, left_border_color="C76E3A")


def _add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BORDER_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_table(doc, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None):
    """Add a branded table. col_widths in inches; None = auto."""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"

    # Set column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            if w is not None:
                for cell in tbl.columns[i].cells:
                    cell.width = Inches(w)

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, FOREST_HEX)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = CREAM

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        bg = CREAM_HEX if r_idx % 2 == 0 else ALT_ROW_HEX
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            _set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = "Calibri"
            run.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def _add_page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def _add_cover(doc: Document):
    # Large top space
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = title.add_run("Evidor")
    r1.font.name = "Calibri"
    r1.font.size = Pt(36)
    r1.font.bold = True
    r1.font.color.rgb = FOREST
    r2 = title.add_run(".ai")
    r2.font.name = "Calibri"
    r2.font.size = Pt(36)
    r2.font.color.rgb = AMBER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rs = sub.add_run("Core Architecture Technical Report")
    rs.font.name = "Calibri"
    rs.font.size = Pt(18)
    rs.font.color.rgb = SLATE

    doc.add_paragraph()

    meta_items = [
        f"Version 1.0",
        f"Region: ap-southeast-1 (Singapore)",
        f"Prepared: {date.today().strftime('%d %B %Y')}",
        "Classification: Confidential — Internal Use Only",
    ]
    for item in meta_items:
        m = doc.add_paragraph()
        r = m.add_run(item)
        r.font.name = "Calibri"
        r.font.size = Pt(10)
        r.font.color.rgb = MUTED
        m.paragraph_format.space_after = Pt(2)

    _add_page_break(doc)


# ---------------------------------------------------------------------------
# Sections
# ---------------------------------------------------------------------------

def _section_overview(doc):
    _add_h1(doc, "1. Project Overview")
    _add_rule(doc)
    _add_para(doc,
        "Evidor.ai is a multi-layer AI message routing system deployed on AWS. "
        "It reads messages from a chatbot interface and dispatches them to the appropriate "
        "AI vendor based on message context, safety assessment, intent classification, and "
        "business rules. Every message passes through five sequential processing layers before "
        "a response is returned to the requester.")
    doc.add_paragraph()
    _add_info_box(doc,
        "Design philosophy: Each layer makes a single focused decision and hands off to the next. "
        "No layer reaches back to a prior one. Fail-open, not fail-closed — "
        "a timeout never hard-blocks a user.")

    _add_h2(doc, "High-Level Request Flow")
    _add_table(doc,
        headers=["Step", "Component", "Decision"],
        rows=[
            ["1", "API Gateway + Lambda Authoriser", "Validate JWT against Cognito JWKS"],
            ["2", "SQS incoming queue", "Buffer and decouple ingestion from processing"],
            ["3", "Orchestrator + Presidio sidecar", "Redact PII — single pass, applied uniformly"],
            ["4", "Bouncer", "Safety gate: rule check + Haiku micro-classifier"],
            ["5", "Intent Classifier", "Classify intent: fast path (embedding) or deep path (Sonnet)"],
            ["6", "Routing Strategist", "Select vendor + apply policy engine"],
            ["7", "Vendor Adapter", "Stream response via Bedrock inference profile"],
            ["8", "Orchestrator (output)", "Restore PII tokens + leak detection + audit"],
            ["9", "WebSocket / SSE", "Return restored chunks to client"],
        ],
        col_widths=[0.5, 2.2, None],
    )
    _add_para(doc,
        "Every request carries a correlation_id set at Orchestrator entry as a Python ContextVar. "
        "Every log line, queue message, and CloudWatch trace emits this ID — enabling end-to-end "
        "tracing across all layers without passing it as a function argument.")


def _section_tech_stack(doc):
    _add_page_break(doc)
    _add_h1(doc, "2. Technology Stack")
    _add_rule(doc)
    _add_table(doc,
        headers=["Category", "Technology", "Notes"],
        rows=[
            ["Language", "Python 3.12", "Consistent across all services and Dockerfiles"],
            ["Web framework", "FastAPI + FastMCP", "Async; MCP server exposure for external agents"],
            ["LLM gateway", "AWS Bedrock (boto3)", "All LLM calls — no direct vendor API clients"],
            ["LLM adapter", "LiteLLM (pinned)", "Vendor adapter layer; avoid 1.82.7/1.82.8 (CVE)"],
            ["Data validation", "Pydantic v2", "All handoff models — no dataclasses, no attrs"],
            ["Dependency manager", "uv", "uv.lock committed; no requirements.txt"],
            ["Queuing", "AWS SQS", "Incoming, escalation, dead-letter queues"],
            ["Caching", "ElastiCache Redis", "User context, session history, PII token vault"],
            ["Config store", "DynamoDB", "Routing rules, audit log"],
            ["Auth", "Amazon Cognito + JWT (RS256)", "Validated at API Gateway edge; never re-validated downstream"],
            ["IaC", "Terraform", "All infra under infra/terraform/"],
            ["Observability", "CloudWatch + X-Ray", "Metrics, traces, S3 archival (365-day retention)"],
            ["PII redaction", "Microsoft Presidio", "ECS Fargate sidecar; custom SG recognisers"],
            ["Testing", "pytest + pytest-asyncio", "fakeredis, LocalStack, moto; 85% coverage gate"],
            ["Region", "ap-southeast-1 (Singapore)", "Hard-coded; data residency enforced at IAM"],
        ],
        col_widths=[1.5, 2.0, None],
    )

    _add_h2(doc, "Bedrock Model IDs (Confirmed Active — ap-southeast-1)")
    _add_para(doc,
        "All model calls use APAC cross-region inference profiles. Raw model IDs fail with "
        "\"on-demand throughput not supported\". Verified via aws bedrock list-inference-profiles.")
    _add_table(doc,
        headers=["Role", "Inference Profile ID"],
        rows=[
            ["Bouncer micro-classifier", "apac.anthropic.claude-3-haiku-20240307-v1:0"],
            ["Intent classifier (deep path)", "apac.anthropic.claude-3-5-sonnet-20241022-v2:0"],
            ["Routing strategist arbitration", "apac.anthropic.claude-3-haiku-20240307-v1:0"],
            ["Latest available (Sonnet 4)", "apac.anthropic.claude-sonnet-4-20250514-v1:0"],
            ["Amazon Nova Pro / Lite / Micro", "apac.amazon.nova-{pro,lite,micro}-v1:0"],
        ],
        col_widths=[2.2, None],
    )


def _section_layers(doc):
    _add_page_break(doc)
    _add_h1(doc, "3. Processing Layers")
    _add_rule(doc)

    # Bouncer
    _add_h2(doc, "Layer 1 — Bouncer")
    _add_para(doc,
        "The first gate. Rejects invalid, malicious, or abusive input before any downstream "
        "LLM cost is incurred. Operates within a 200 ms total budget.")
    _add_table(doc,
        headers=["Sub-layer", "Mechanism", "Latency", "Result on failure"],
        rows=[
            ["Rule gate",
             "Length check, prompt injection regex, banned user (Redis), rate limit (Redis)",
             "< 1 ms",
             "allowed=False — hard block"],
            ["Haiku micro-classifier",
             "Bedrock Haiku (max 50 tokens). Returns {pass, reason, confidence}.",
             "~50–150 ms",
             "Timeout/error → fail open (allowed=True, timed_out=True)"],
        ],
        col_widths=[1.2, 2.8, 0.9, None],
    )
    _add_amber_box(doc,
        "Critical design invariant — Fail-Open: A Bouncer timeout or Bedrock error always "
        "resolves to allowed=True. The Haiku LLM never hard-blocks a user — only the rule gate "
        "does. A CloudWatch metric BouncerTimeout is incremented on every fail-open event.")
    doc.add_paragraph()

    # Classifier
    _add_h2(doc, "Layer 2 — Intent Classifier")
    _add_para(doc,
        "Decides what the user wants so the Strategist knows who should handle it. "
        "Two execution paths selected by message complexity.")
    _add_table(doc,
        headers=["Path", "Mechanism", "Latency", "Threshold"],
        rows=[
            ["Fast path",
             "Embedding similarity (Bedrock Titan) against known intent vectors",
             "~20 ms",
             "Confidence >= 0.82"],
            ["Deep path",
             "Claude Sonnet with MCP tool calls (taxonomy, session history, guardrails, entity context)",
             "~300–800 ms",
             "Confidence < 0.82"],
        ],
        col_widths=[1.0, 3.2, 0.9, None],
    )
    _add_h3(doc, "Intent Taxonomy")
    _add_bullet(doc, "code_assistance — routed to Claude Sonnet")
    _add_bullet(doc, "general_qa — routed to Claude Sonnet")
    _add_bullet(doc, "simple_transactional — routed to Claude Haiku (cost-optimised)")
    _add_bullet(doc, "out_of_scope / ambiguous — escalated to human review queue")
    _add_h3(doc, "MCP Tools (deep path — called in parallel via asyncio.gather)")
    _add_bullet(doc, "get_intent_taxonomy — taxonomy from DynamoDB")
    _add_bullet(doc, "get_session_history — last 3 turns for pronoun resolution")
    _add_bullet(doc, "check_guardrails — Bedrock Guardrails content check")
    _add_bullet(doc, "get_entity_context — known entities for this user")
    _add_para(doc, "Escalates to human review if confidence < 0.6. Never guesses a vendor when uncertain.")
    doc.add_paragraph()

    # Strategist
    _add_h2(doc, "Layer 3 — Routing Strategist")
    _add_para(doc,
        "Selects which vendor handles the request, constructs the fallback chain, "
        "and runs the policy engine to enforce data residency and compliance rules.")
    _add_table(doc,
        headers=["Confidence", "Execution path"],
        rows=[
            [">= 0.85", "Deterministic rule lookup in DynamoDB — no LLM call"],
            ["0.50 – 0.85", "Haiku arbitration call (max 80 tokens)"],
            ["< 0.50", "Escalate to human review queue"],
        ],
        col_widths=[1.4, None],
    )
    _add_h3(doc, "Policy Engine (runs after vendor selection, before output)")
    _add_bullet(doc, "Data residency: SG users must be served via Bedrock ap-southeast-1 only")
    _add_bullet(doc, "Legal/compliance blocks: MAS-regulated queries, tenancy law triggers")
    _add_bullet(doc, "Cross-jurisdiction flags: overseas content for SG-resident user")
    doc.add_paragraph()

    # Adapters
    _add_h2(doc, "Layer 4 — Vendor Adapters")
    _add_para(doc,
        "All vendors are unified behind a single Bedrock interface. "
        "A single boto3 client handles Claude, Llama, and Mistral models. "
        "Streaming is enabled by default for a chat-like user experience. "
        "LiteLLM is used as the adapter layer — it handles multi-model API differences, "
        "fallbacks, retries, and cost tracking.")
    _add_info_box(doc,
        "Bedrock-only constraint: No direct OpenAI, Anthropic, or Google API clients are used "
        "anywhere in the codebase. CI enforces this with a forbidden-import check. All vendor "
        "calls go through boto3.client(\"bedrock-runtime\") or LiteLLM configured against Bedrock.")


def _section_compute(doc):
    _add_page_break(doc)
    _add_h1(doc, "4. Compute Architecture")
    _add_rule(doc)
    _add_para(doc,
        "Phase 1 deployment: every layer is one ECS Fargate service. "
        "Stateless/bursty components run on Lambda. "
        "ECS is used where persistent connections or tight latency budgets prohibit Lambda cold starts.")
    _add_table(doc,
        headers=["Component", "Platform", "Rationale"],
        rows=[
            ["API Gateway authoriser", "Lambda", "Stateless JWT validation; millisecond execution"],
            ["SQS consumers", "Lambda", "Native SQS trigger; scales with queue depth"],
            ["Rule-based gate", "Lambda", "Pure Python, stateless, bursty"],
            ["Routing strategist", "Lambda", "80–90% pure rule lookups; occasional LLM well within 15 min limit"],
            ["Orchestrator", "Lambda (VPC-attached)", "Stateless SQS consumer; VPC needed for Redis + Presidio access"],
            ["Haiku classifier (Bouncer)", "ECS Fargate", "200 ms budget cannot absorb Lambda cold start"],
            ["Intent classifier", "ECS Fargate", "Persistent MCP clients; in-memory taxonomy cache"],
            ["Vendor adapters", "ECS Fargate", "Streaming responses require persistent HTTP connections"],
            ["WebSocket server", "ECS Fargate", "Lambda cannot hold persistent WebSocket connections"],
            ["Presidio sidecar", "ECS Fargate", "~200 MB spaCy model must stay resident between requests"],
            ["Admin dashboard (FastAPI)", "ECS Fargate", "Persistent Redis + DynamoDB connections; low concurrency"],
        ],
        col_widths=[2.0, 1.6, None],
    )


def _section_redaction(doc):
    _add_page_break(doc)
    _add_h1(doc, "5. PII Redaction & Data Privacy")
    _add_rule(doc)
    _add_para(doc,
        "Single redaction pass at Orchestrator entry, applied uniformly to every LLM call. "
        "The same redacted message flows through Haiku (Bouncer), Sonnet (Classifier), "
        "and the vendor adapter. No per-layer redaction. No skipping redaction for any layer.")

    _add_h2(doc, "What Gets Redacted (High-Sensitivity)")
    _add_table(doc,
        headers=["Entity Type", "Description"],
        rows=[
            ["SG_NRIC / SG_FIN", "Singapore National Registration Identity Card / Foreign ID Number"],
            ["SG_UEN", "Unique Entity Number (businesses)"],
            ["SG_PASSPORT", "Singapore passport numbers"],
            ["CREDIT_CARD / IBAN_CODE", "Payment card and bank account numbers"],
            ["Large currency figures", "Amounts exceeding configured threshold"],
        ],
        col_widths=[2.0, None],
    )
    _add_para(doc,
        "Contextual entities (LOCATION, PERSON, DATE_TIME, PHONE_NUMBER, EMAIL_ADDRESS) "
        "remain in the message — they are needed for accurate intent classification.")

    _add_h2(doc, "Redaction Pipeline")
    _add_bullet(doc, "Input redaction via Microsoft Presidio (ECS Fargate sidecar) with custom Singapore recognisers")
    _add_bullet(doc, "Tokens stored in Redis vault as vault:{correlation_id}:{token} with 5-minute TTL")
    _add_bullet(doc, "All three LLM calls see the same redacted message")
    _add_bullet(doc, "Output restoration unwinds tokens after vendor response")
    _add_bullet(doc, "Leak detector scans restored response — strips unexpected PII before delivery")
    _add_bullet(doc, "Audit log records entity types and counts only — never values")
    _add_bullet(doc, "Vault keys eagerly deleted in finally block after each request")

    _add_h2(doc, "Streaming Redaction")
    _add_para(doc,
        "Buffer-and-scan pattern: 200-character buffer with a 50-character safety margin "
        "to handle mid-entity splits across chunk boundaries. Constants defined in "
        "redactor/streaming_redactor.py.")
    _add_info_box(doc,
        "No PII in logs: Allowlist-based structured logging (structlog) is used throughout. "
        "Only fields on the allowlist are emitted. Errors are logged as type names only — never "
        "as messages, which may contain echoed user input. Bedrock invocation logging is "
        "explicitly disabled in Terraform.")


def _section_auth(doc):
    _add_page_break(doc)
    _add_h1(doc, "6. Authentication & Security")
    _add_rule(doc)

    _add_h2(doc, "Authentication Flow")
    _add_bullet(doc, "Client authenticates against Amazon Cognito User Pool and receives a JWT access token (RS256)")
    _add_bullet(doc, "Client sends API request with Authorization: Bearer <token>")
    _add_bullet(doc, "API Gateway Lambda authoriser validates JWT signature against Cognito JWKS and checks exp, iss, aud, token_use claims")
    _add_bullet(doc, "Invalid tokens rejected at the edge with HTTP 401 — no payload reaches SQS or the pipeline")
    _add_bullet(doc, "Validated user_sub (Cognito UUID) injected into request context and flows through all layers")

    _add_table(doc,
        headers=["Token", "TTL", "Notes"],
        rows=[
            ["Access token", "60 minutes", "Sent with every API request"],
            ["Refresh token", "30 days", "Never sent to the API; client-side only"],
            ["PII vault token", "5 minutes", "Redis key TTL; eager delete after request completes"],
        ],
        col_widths=[1.4, 1.2, None],
    )

    _add_h2(doc, "IAM & Data Residency")
    _add_bullet(doc, "Bedrock IAM uses APAC cross-region inference profile IDs — data residency enforced by profile selection, not IAM region conditions")
    _add_bullet(doc, "Admin service IAM role explicitly denies bedrock:*, sqs:SendMessage on the incoming queue, and dynamodb:DeleteItem")
    _add_bullet(doc, "Admin dashboard sits behind a separate ALB listener rule (/admin/*) with IP allowlisting")
    _add_bullet(doc, "Raw message hash is stored for audit; raw message itself is never persisted or propagated beyond the Orchestrator")

    _add_h2(doc, "Key Security Principles")
    _add_table(doc,
        headers=["#", "Principle"],
        rows=[
            ["1", "All LLM calls go through Bedrock — no direct vendor API clients"],
            ["2", "JWT validated at the edge — pipeline code never re-validates"],
            ["3", "Single redaction policy applied uniformly to every LLM call"],
            ["4", "Fail-open Bouncer — timeout never hard-blocks a user"],
            ["5", "Correlation ID threads everything via ContextVar — never passed as an argument"],
            ["6", "No PII in logs ever — allowlist-based structured logging, error type not error message"],
            ["7", "Bedrock invocation logging disabled — application-level audit is the source of truth"],
            ["8", "Escalate, don't guess — low confidence routes to human review, never picks a default vendor"],
            ["9", "Data residency enforced at APAC inference profile selection, not IAM region conditions"],
            ["10", "Raw message never leaves the Orchestrator — only redacted_message in PipelineEnvelope"],
        ],
        col_widths=[0.35, None],
    )


def _section_admin(doc):
    _add_page_break(doc)
    _add_h1(doc, "7. Admin Dashboard")
    _add_rule(doc)
    _add_para(doc,
        "React SPA served from S3 + CloudFront. Backend is a dedicated FastAPI service on "
        "ECS Fargate. Read-heavy design — the only write surfaces are the routing rule editor, "
        "escalation queue actions, and tier overrides.")
    _add_table(doc,
        headers=["Deployment detail", "Value"],
        rows=[
            ["Frontend hosting", "S3 bucket + CloudFront distribution (E39VIJ6ZMJDE2V)"],
            ["URL", "https://d13ux0ouuro9ex.cloudfront.net"],
            ["Backend", "FastAPI on ECS Fargate, behind ALB /admin/* listener rule"],
            ["Auth", "Same Cognito User Pool as the main pipeline (JWT RS256)"],
            ["CloudWatch namespace", "AIRouter/Admin"],
            ["Alarm thresholds", "Escalation queue > 50 messages OR error rate > 5% for 5 min"],
        ],
        col_widths=[2.0, None],
    )
    _add_h2(doc, "Dashboard Views")
    _add_bullet(doc, "Pipeline Health — per-layer throughput, latency (p50/p99/p999), error rate, escalation rate")
    _add_bullet(doc, "Bouncer — pass/fail/escalate rates, confidence histogram, top blocked patterns")
    _add_bullet(doc, "Classifier — intent distribution, fast vs deep path split, confidence distribution")
    _add_bullet(doc, "Strategist — vendor selection breakdown, policy engine trigger counts")
    _add_bullet(doc, "Escalation Queue — pending SQS human review messages (redacted previews only); approve, reject, or requeue")
    _add_bullet(doc, "Routing Rules — DynamoDB CRUD for intent→vendor mapping; changes take effect immediately")
    _add_bullet(doc, "Audit Log — entity type counts, policies applied, vendor used per correlation_id; S3-backed, 365-day retention")
    _add_info_box(doc,
        "Test Console: Embedded tool allowing admins to submit a message and trace it through the full "
        "pipeline (or dry-run stopping before vendor invocation). Shows BounceResult → ClassifiedIntent "
        "→ RoutingPlan → vendor response with per-layer latencies. Chat history persisted in localStorage "
        "(capped at 100 entries). All traces log redacted messages only.")


def _section_handoffs(doc):
    _add_page_break(doc)
    _add_h1(doc, "8. Key Handoff Contracts")
    _add_rule(doc)
    _add_para(doc,
        "All inter-layer contracts are defined as Pydantic v2 models in shared/models.py. "
        "Every model round-trips through JSON in contract tests.")
    _add_table(doc,
        headers=["Model", "Key fields"],
        rows=[
            ["PipelineEnvelope",
             "correlation_id, user_sub, session_id, redacted_message, raw_message_hash, "
             "entity_types_redacted, entity_count, was_redacted, timestamp, bedrock_region, source_ip"],
            ["BounceResult", "allowed, reason, layer, confidence, escalate, timed_out"],
            ["ClassifiedIntent",
             "intent, sub_intent, domain, confidence, entities, resolved_message, multi_domain, escalate, reasoning"],
            ["RoutingPlan",
             "primary_vendor, fallback_chain, context, applied_policies, policy_modified, blocked"],
            ["RedactionResult",
             "redacted_message, entity_types_found, entity_count, was_redacted, correlation_id"],
        ],
        col_widths=[1.7, None],
    )
    _add_h2(doc, "MCP Architecture")
    _add_para(doc,
        "The router is exposed as a FastMCP server (ECS Fargate) so external agents can call "
        "route_message and get_routing_status tools.")
    _add_table(doc,
        headers=["Layer", "Outbound MCP tools"],
        rows=[
            ["Bouncer", "None — rule gate only"],
            ["Classifier", "get_intent_taxonomy, get_session_history, check_guardrails, get_entity_context"],
            ["Strategist", "check_vendor_health, get_routing_rules"],
        ],
        col_widths=[1.4, None],
    )
    _add_amber_box(doc,
        "MCP tool calls in the Classifier deep path are issued in parallel via asyncio.gather — "
        "sequential calls would blow the latency budget. Each MCP server owns one domain of data: "
        "replaceable, testable, and reusable.")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def build():
    out_path = Path("reports/architecture-report.docx")
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configure_styles(doc)

    _add_cover(doc)
    _section_overview(doc)
    _section_tech_stack(doc)
    _section_layers(doc)
    _section_compute(doc)
    _section_redaction(doc)
    _section_auth(doc)
    _section_admin(doc)
    _section_handoffs(doc)

    doc.save(str(out_path))
    print(f"Report written to: {out_path.resolve()}")


if __name__ == "__main__":
    build()
